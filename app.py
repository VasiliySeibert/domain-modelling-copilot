from flask import Flask, render_template, request, jsonify, session
from openai_client import OpenAIClient  # Utility class
from gpt2 import gpt_v2_interface
from docx import Document
from flask import send_from_directory
import os

app = Flask(__name__)
app.secret_key = "supersecretkey"  # Required for session handling

# Initialize the OpenAI client once
OpenAIClient.initialize()

chat_history = []
scenario = []  # Stores all scenarios generated during the chat session

@app.route("/")
def home():
    global chat_history
    chat_history = []  # Reset chat history on page load
    return render_template("index.html")

# Function to handle chat messages
@app.route("/chat", methods=["POST"])
def chat():
    """Handle chat messages and classify them as 'general' or 'scenario'."""
    global chat_history, scenario
    try:
        # Get user message
        user_message = request.json.get("message", "").strip()
        if not user_message:
            return jsonify({"error": "Message is required"}), 400

        # Retrieve the user's name from the session
        user_name = session.get("user_name")
        if not user_name:
            return jsonify({"error": "User name is not set. Please submit your name first."}), 400

        # Add the user's message to the chat history
        chat_history.append({"role": "user", "content": user_message})

        # Classify the input as "general" or "scenario"
        classification = classify_input(user_message)

        if classification == "general":
            # Generate a general response
            response = generate_general(user_name, chat_history)
            bot_reply = response.json["response"]
            chat_history.append({"role": "assistant", "content": bot_reply})
            return response

        elif classification == "scenario":
            # Generate the detailed scenario
            detailed_description = generate_scenario(user_message)

            # Generate the summary from the detailed scenario
            summary = generate_summary(detailed_description)
            scenarios.append(detailed_description)
            filename = f"{user_name}_scenario_{len(scenarios)}.docx"
            save_scenario_to_word(detailed_description, filename)
            chat_history.append({"role": "assistant", "content": summary})

            # Return both the detailed scenario and the summary
            return jsonify({"scenario": detailed_description, "summary": summary})

        else:
            # Invalid classification
            return jsonify({"error": "Unable to classify the input. Please rephrase your query."}), 400

    except Exception as e:
        print(f"Error in /chat endpoint: {e}")
        return jsonify({"error": "An error occurred"}), 500

# Function to handle the name submission
@app.route("/submit_name", methods=["POST"])
def submit_name():
    """Store the user's name in the session."""
    try:
        user_name = request.json.get("name", "").strip()
        if not user_name:
            return jsonify({"error": "Name is required"}), 400

        # Store the name in the session
        session["user_name"] = user_name
        return jsonify({"message": "Name saved successfully!", "name": user_name})
    except Exception as e:
        print(f"Error storing name: {e}")
        return jsonify({"error": "An error occurred while storing the name"}), 500

# Function to generate a general response using GPT
def generate_general(user_name, chat_history):
    """Handle general input and return a response using GPT."""
    try:
        # Prepare the prompts with the chat history
        prompts = [
            {"role": "system", "content": f"You are a helpful assistant. The user's name is {user_name}."}
        ] + chat_history

        # Get the OpenAI client
        client = OpenAIClient.get_client()

        # Get response from GPT
        response = client.chat.completions.create(
            model=os.getenv("GPT_MODEL"), messages=prompts
        )
        bot_reply = response.choices[0].message.content.strip()

        # Add the bot's response to the chat history
        chat_history.append({"role": "assistant", "content": bot_reply})

        # Return the raw bot reply
        return jsonify({"response": bot_reply, "history": chat_history})
    except Exception as e:
        print(f"Error genreating general response: {e}")
        return jsonify({"error": "An error occurred while processing your request."}), 500
    

def generate_scenario(scenario_text=None):
    """Generate a detailed scenario from the given user input."""
    try:
        # If scenario_text is not provided, get it from the request (for direct API calls)
        if scenario_text is None:
            scenario_text = request.json.get("message", "").strip()
            if not scenario_text:
                return jsonify({"error": "Scenario text is required"}), 400

        # Use LLM to generate a detailed scenario description
        prompts = [
            {"role": "system", "content": "You are an expert in converting user inputs into detailed scenarios. Only respond with the detailed scenario, do not add anything else like title, conclusion, etc."},
            {"role": "user", "content": f"Generate a detailed scenario for the following input:\n\n{scenario_text}"}
        ]

        # Get the OpenAI client
        client = OpenAIClient.get_client()

        # Call the GPT model to generate the detailed scenario
        response = client.chat.completions.create(
            model=os.getenv("GPT_MODEL"),
            messages=prompts
        )
        detailed_description = response.choices[0].message.content.strip()

        # Return the detailed scenario
        return detailed_description
    except Exception as e:
        print(f"Error generating scenario: {e}")
        return "An error occurred while generating the scenario."

@app.route("/generate_uml", methods=["POST"])
def generate_uml():
    """Generate UML text from the given scenario text."""
    try:
        # Get the scenario text from the request
        scenario_text = request.json.get("scenarioText", "").strip()
        if not scenario_text:
            return jsonify({"error": "Scenario text is required"}), 400

        # Get the OpenAI client
        client = OpenAIClient.get_client()

        # Call gpt_v2_interface to process the scenario and generate PlantUML
        plant_uml = gpt_v2_interface(scenario_text, client)

        # Return the generated PlantUML
        return jsonify({"plantuml": plant_uml})
    except Exception as e:
        print(f"Error generating UML: {e}")
        return jsonify({"error": "An error occurred while generating the UML"}), 500

@app.route("/generate_summary", methods=["POST"])
def generate_summary():
    """Generate a summary from the given detailed scenario."""
    try:
        # Get the detailed scenario from the request
        detailed_description = request.json.get("detailed_description", "").strip()
        if not detailed_description:
            return jsonify({"error": "Detailed description is required"}), 400

        # Use GPT to generate a summary
        summary = generate_summary(detailed_description)

        # Return the summary
        return jsonify({"summary": summary})
    except Exception as e:
        print(f"Error generating summary: {e}")
        return jsonify({"error": "An error occurred while generating the summary"}), 500

def generate_summary(detailed_description):
    """Generate a summary from the given detailed scenario."""
    try:
        # Use GPT to generate a summary
        prompts = [
            {"role": "system", "content": "You are an expert in summarizing detailed scenarios into concise summaries. A summary should capture the essence of the scenario in one or two sentences."},
            {"role": "user", "content": f"Summarize the following scenario in one or two sentences:\n\n{detailed_description}"}
        ]

        # Get the OpenAI client
        client = OpenAIClient.get_client()

        # Call the GPT model to generate the summary
        response = client.chat.completions.create(
            model=os.getenv("GPT_MODEL"),
            messages=prompts
        )
        summary = response.choices[0].message.content.strip()

        # Return the summary
        return summary
    except Exception as e:
        print(f"Error generating summary: {e}")
        return "An error occurred while generating the summary."

def save_scenario_to_word(scenario_text, filename="scenario.docx"):
    project_dir = os.path.join(os.getcwd(), "project")
    os.makedirs(project_dir, exist_ok=True)

    file_path = os.path.join(project_dir, filename)
    doc = Document()
    doc.add_heading("Generated Scenario", 0)
    doc.add_paragraph(scenario_text)
    doc.save(file_path)
    return file_path

@app.route("/list_word_files", methods=["GET"])
def list_word_files():
    project_dir = os.path.join(os.getcwd(), "project")
    if not os.path.exists(project_dir):
        return jsonify({"files": []})
    files = [f for f in os.listdir(project_dir) if f.endswith(".docx")]
    return jsonify({"files": files})

@app.route("/download_word/<filename>")
def download_word(filename):
    project_dir = os.path.join(os.getcwd(), "project")
    return send_from_directory(directory=project_dir, path=filename, as_attachment=True)

@app.route("/read_word/<filename>", methods=["GET"])
def read_word(filename):
    try:
        project_dir = os.path.join(os.getcwd(), "project")
        file_path = os.path.join(project_dir, filename)

        if not os.path.exists(file_path):
            return jsonify({"error": "File not found"}), 404

        doc = Document(file_path)
        full_text = "\n".join([para.text for para in doc.paragraphs])
        return jsonify({"content": full_text})
    except Exception as e:
        print(f"Error reading Word file: {e}")
        return jsonify({"error": "Failed to read file"}), 500

def classify_input(user_message):
    """Classify the user input as 'general' or 'scenario'."""
    try:
        # Get the OpenAI client
        client = OpenAIClient.get_client()

        # Prepare the classification prompt
        prompts = [
            {"role": "system", "content": "You are an expert assistant trained to classify user inputs. Response with one of the following: 'general' or 'scenario'."},
            {"role": "user", "content": f"Classify the following input as either 'general' or 'scenario'.).\n\nInput: {user_message}"}
        ]

        # Get the classification response
        response = client.chat.completions.create(
            model=os.getenv("GPT_MODEL"),
            messages=prompts
        )
        classification = response.choices[0].message.content.strip().lower()
        return classification
    except Exception as e:
        print(f"Error classifying input: {e}")
        return "general"  # Default to "general" in case of an error

@app.route("/get_scenarios", methods=["GET"])
def get_scenarios():
    """Retrieve all stored scenarios."""
    global scenario
    return jsonify({"scenarios": scenario})

if __name__ == "__main__":
    app.run(debug=True)